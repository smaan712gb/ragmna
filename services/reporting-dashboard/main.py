"""
Reporting Dashboard Service
Creates comprehensive reports and interactive dashboards for M&A analysis
Generates Word documents, PDF reports, and dashboard data
"""

import os
import json
import logging
from flask import Flask, request, jsonify, send_file
from functools import wraps
from typing import Dict, Any, List, Optional
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import base64

app = Flask(__name__)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Service configurations
SERVICE_API_KEY = os.getenv('SERVICE_API_KEY')

class ReportGenerator:
    """Comprehensive report and dashboard generator"""

    def __init__(self):
        self.styles = {
            'title': {'size': 16, 'bold': True},
            'heading': {'size': 14, 'bold': True},
            'subheading': {'size': 12, 'bold': True},
            'normal': {'size': 11},
            'caption': {'size': 10, 'italic': True}
        }

    def generate_comprehensive_report(self, analysis_data: Dict[str, Any]) -> bytes:
        """Generate comprehensive Word document report"""

        doc = Document()
        self._add_title_page(doc, analysis_data)
        self._add_executive_summary(doc, analysis_data)
        self._add_company_overview(doc, analysis_data)
        self._add_financial_analysis(doc, analysis_data)
        self._add_valuation_analysis(doc, analysis_data)
        self._add_due_diligence_findings(doc, analysis_data)
        self._add_recommendations(doc, analysis_data)
        self._add_appendices(doc, analysis_data)

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return output.getvalue()

    def generate_dashboard_data(self, analysis_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate interactive dashboard data"""

        dashboard_data = {
            'summary_metrics': self._generate_summary_metrics(analysis_data),
            'charts': self._generate_chart_data(analysis_data),
            'valuation_comparison': self._generate_valuation_comparison(analysis_data),
            'risk_heatmap': self._generate_risk_heatmap(analysis_data),
            'timeline_data': self._generate_timeline_data(analysis_data),
            'generated_at': datetime.now().isoformat()
        }

        return dashboard_data

    def _add_title_page(self, doc: Document, data: Dict[str, Any]):
        """Add title page to document"""

        title = doc.add_heading('M&A Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add company information
        target_symbol = data.get('target_symbol', 'Unknown')
        acquirer_symbol = data.get('acquirer_symbol', 'N/A')

        doc.add_paragraph(f"Target Company: {target_symbol}")
        if acquirer_symbol != 'N/A':
            doc.add_paragraph(f"Acquirer Company: {acquirer_symbol}")

        doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%B %d, %Y')}")

        # Add page break
        doc.add_page_break()

    def _add_executive_summary(self, doc: Document, data: Dict[str, Any]):
        """Add executive summary section"""

        doc.add_heading('Executive Summary', 1)

        # Key findings
        final_report = data.get('final_report', {})
        summary = final_report.get('summary', {})

        doc.add_paragraph("Key Findings:", style='List Bullet')
        doc.add_paragraph(f"• Company Classification: {summary.get('company_classification', 'Unknown').title()}", style='List Bullet')
        doc.add_paragraph(f"• Overall Risk Assessment: {summary.get('risk_assessment', 'Unknown').title()}", style='List Bullet')

        # Valuation summary
        valuation = data.get('valuation', {})
        dcf = valuation.get('dcf', {})
        cca = valuation.get('cca', {})

        if dcf.get('final_valuation'):
            dcf_value = dcf['final_valuation'].get('equity_value_per_share', 0)
            doc.add_paragraph(f"• DCF Valuation: ${dcf_value:,.2f} per share", style='List Bullet')

        if cca.get('implied_valuation', {}).get('blended_valuation'):
            cca_value = cca['implied_valuation']['blended_valuation'].get('blended_price_per_share', 0)
            doc.add_paragraph(f"• Comparable Analysis: ${cca_value:,.2f} per share", style='List Bullet')

        # Recommendations
        recommendations = summary.get('recommendations', [])
        if recommendations:
            doc.add_paragraph("Key Recommendations:", style='List Bullet')
            for rec in recommendations[:3]:
                doc.add_paragraph(f"• {rec}", style='List Bullet')

    def _add_company_overview(self, doc: Document, data: Dict[str, Any]):
        """Add company overview section"""

        doc.add_heading('Company Overview', 1)

        target_data = data.get('target_data', {})
        profile = target_data.get('profile', [{}])[0] if target_data.get('profile') else {}
        market = target_data.get('market', {})

        # Basic information table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'

        cells = [
            ('Company Name', profile.get('companyName', 'N/A')),
            ('Symbol', data.get('target_symbol', 'N/A')),
            ('Sector', profile.get('sector', 'N/A')),
            ('Industry', profile.get('industry', 'N/A')),
            ('Market Cap', f"${market.get('marketCap', 0):,.0f}"),
            ('Current Price', f"${market.get('price', 0):,.2f}")
        ]

        for i, (label, value) in enumerate(cells):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value

        # Classification
        classification = data.get('classification', {})
        doc.add_paragraph(f"Company Classification: {classification.get('primary_classification', 'Unknown').title()}")

    def _add_financial_analysis(self, doc: Document, data: Dict[str, Any]):
        """Add financial analysis section"""

        doc.add_heading('Financial Analysis', 1)

        financial_model = data.get('financial_model', {})

        # Income statement projections
        doc.add_heading('Income Statement Projections', 2)

        income_stmt = financial_model.get('income_statement', [])
        if income_stmt:
            table = doc.add_table(rows=len(income_stmt) + 1, cols=5)
            table.style = 'Table Grid'

            # Headers
            headers = ['Year', 'Revenue', 'EBIT', 'Net Income', 'EPS']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header

            # Data
            for row, stmt in enumerate(income_stmt, 1):
                table.cell(row, 0).text = str(stmt.get('year', ''))
                table.cell(row, 1).text = f"${stmt.get('revenue', 0):,.0f}"
                table.cell(row, 2).text = f"${stmt.get('operating_income', 0):,.0f}"
                table.cell(row, 3).text = f"${stmt.get('net_income', 0):,.0f}"
                table.cell(row, 4).text = f"${stmt.get('eps', 0):.2f}"

    def _add_valuation_analysis(self, doc: Document, data: Dict[str, Any]):
        """Add valuation analysis section"""

        doc.add_heading('Valuation Analysis', 1)

        valuation = data.get('valuation', {})

        # DCF Analysis
        doc.add_heading('Discounted Cash Flow Analysis', 2)

        dcf = valuation.get('dcf', {})
        if dcf.get('final_valuation'):
            fv = dcf['final_valuation']
            doc.add_paragraph(f"Enterprise Value: ${fv.get('enterprise_value', 0):,.0f}")
            doc.add_paragraph(f"Equity Value: ${fv.get('equity_value', 0):,.0f}")
            doc.add_paragraph(f"Price per Share: ${fv.get('equity_value_per_share', 0):,.2f}")
            doc.add_paragraph(f"WACC: {dcf.get('wacc', 0):.1%}")

        # Comparable Analysis
        doc.add_heading('Comparable Company Analysis', 2)

        cca = valuation.get('cca', {})
        if cca.get('implied_valuation', {}).get('blended_valuation'):
            bv = cca['implied_valuation']['blended_valuation']
            doc.add_paragraph(f"Implied Price per Share: ${bv.get('blended_price_per_share', 0):,.2f}")

    def _add_due_diligence_findings(self, doc: Document, data: Dict[str, Any]):
        """Add due diligence findings section"""

        doc.add_heading('Due Diligence Findings', 1)

        dd_report = data.get('due_diligence', {})
        overall_assessment = dd_report.get('overall_assessment', {})

        doc.add_paragraph(f"Overall Risk Level: {overall_assessment.get('overall_risk_level', 'Unknown').title()}")

        # Risk categories
        risk_categories = ['legal_analysis', 'financial_analysis', 'operational_analysis', 'strategic_analysis', 'reputational_analysis']

        for category in risk_categories:
            analysis = dd_report.get(category, {})
            risk_level = analysis.get(f"overall_{category.split('_')[0]}_score", 0)
            if risk_level >= 2.5:
                doc.add_paragraph(f"• {category.replace('_', ' ').title()}: High Risk", style='List Bullet')

    def _add_recommendations(self, doc: Document, data: Dict[str, Any]):
        """Add recommendations section"""

        doc.add_heading('Recommendations', 1)

        final_report = data.get('final_report', {})
        summary = final_report.get('summary', {})
        recommendations = summary.get('recommendations', [])

        for rec in recommendations:
            doc.add_paragraph(f"• {rec}", style='List Bullet')

    def _add_appendices(self, doc: Document, data: Dict[str, Any]):
        """Add appendices with detailed data"""

        doc.add_heading('Appendices', 1)

        # Add detailed financials
        doc.add_heading('Appendix A: Detailed Financial Projections', 2)

        # Add methodology
        doc.add_heading('Appendix B: Valuation Methodology', 2)
        doc.add_paragraph("This analysis uses a combination of DCF and comparable company analysis methods, adjusted for company-specific risk factors and market conditions.")

    def _generate_summary_metrics(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate summary metrics for dashboard"""

        target_data = data.get('target_data', {})
        market = target_data.get('market', {})

        valuation = data.get('valuation', {})
        dcf = valuation.get('dcf', {})
        cca = valuation.get('cca', {})

        return {
            'company_name': target_data.get('profile', [{}])[0].get('companyName', 'Unknown'),
            'symbol': data.get('target_symbol', 'Unknown'),
            'current_price': market.get('price', 0),
            'market_cap': market.get('marketCap', 0),
            'dcf_value': dcf.get('final_valuation', {}).get('equity_value_per_share', 0),
            'cca_value': cca.get('implied_valuation', {}).get('blended_valuation', {}).get('blended_price_per_share', 0),
            'classification': data.get('classification', {}).get('primary_classification', 'unknown'),
            'risk_level': data.get('due_diligence', {}).get('overall_assessment', {}).get('overall_risk_level', 'unknown')
        }

    def _generate_chart_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate chart data for dashboard"""

        charts = {}

        # Valuation comparison chart
        valuation = data.get('valuation', {})
        dcf_value = valuation.get('dcf', {}).get('final_valuation', {}).get('equity_value_per_share', 0)
        cca_value = valuation.get('cca', {}).get('implied_valuation', {}).get('blended_valuation', {}).get('blended_price_per_share', 0)
        current_price = data.get('target_data', {}).get('market', {}).get('price', 0)

        charts['valuation_comparison'] = {
            'labels': ['Current Price', 'DCF Value', 'CCA Value'],
            'values': [current_price, dcf_value, cca_value],
            'type': 'bar'
        }

        # Financial projections chart
        financial_model = data.get('financial_model', {})
        income_stmt = financial_model.get('income_statement', [])

        if income_stmt:
            years = [str(stmt.get('year', '')) for stmt in income_stmt]
            revenues = [stmt.get('revenue', 0) for stmt in income_stmt]
            net_income = [stmt.get('net_income', 0) for stmt in income_stmt]

            charts['revenue_projection'] = {
                'labels': years,
                'values': revenues,
                'type': 'line'
            }

            charts['earnings_projection'] = {
                'labels': years,
                'values': net_income,
                'type': 'line'
            }

        return charts

    def _generate_valuation_comparison(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate valuation comparison data"""

        valuation = data.get('valuation', {})
        current_price = data.get('target_data', {}).get('market', {}).get('price', 0)

        comparisons = []

        # DCF comparison
        dcf = valuation.get('dcf', {})
        if dcf.get('final_valuation'):
            dcf_value = dcf['final_valuation'].get('equity_value_per_share', 0)
            comparisons.append({
                'method': 'DCF',
                'value': dcf_value,
                'premium_discount': (dcf_value - current_price) / current_price if current_price > 0 else 0
            })

        # CCA comparison
        cca = valuation.get('cca', {})
        if cca.get('implied_valuation', {}).get('blended_valuation'):
            cca_value = cca['implied_valuation']['blended_valuation'].get('blended_price_per_share', 0)
            comparisons.append({
                'method': 'Comparable Analysis',
                'value': cca_value,
                'premium_discount': (cca_value - current_price) / current_price if current_price > 0 else 0
            })

        return {'comparisons': comparisons}

    def _generate_risk_heatmap(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate risk heatmap data"""

        dd_report = data.get('due_diligence', {})
        risk_categories = ['legal_analysis', 'financial_analysis', 'operational_analysis', 'strategic_analysis', 'reputational_analysis']

        heatmap_data = []

        for category in risk_categories:
            analysis = dd_report.get(category, {})
            risk_score = analysis.get(f"overall_{category.split('_')[0]}_score", 0)
            heatmap_data.append({
                'category': category.replace('_', ' ').title(),
                'risk_score': risk_score,
                'risk_level': 'High' if risk_score >= 2.5 else 'Medium' if risk_score >= 1.5 else 'Low'
            })

        return {'heatmap_data': heatmap_data}

    def _generate_timeline_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate timeline data for dashboard"""

        financial_model = data.get('financial_model', {})
        income_stmt = financial_model.get('income_statement', [])

        timeline = []

        for stmt in income_stmt:
            timeline.append({
                'year': stmt.get('year', ''),
                'revenue': stmt.get('revenue', 0),
                'net_income': stmt.get('net_income', 0),
                'eps': stmt.get('eps', 0)
            })

        return {'timeline': timeline}

# Global report generator instance
report_generator = ReportGenerator()

def require_api_key(f):
    """Decorator to require API key"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        api_key = request.headers.get('X-API-Key')
        if not api_key or api_key != SERVICE_API_KEY:
            return jsonify({'error': 'Invalid API key'}), 401
        return f(*args, **kwargs)
    return decorated_function

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'service': 'reporting-dashboard',
        'version': '1.0.0'
    })

@app.route('/dashboard/generate', methods=['POST'])
@require_api_key
def generate_dashboard():
    """Generate interactive dashboard data"""
    try:
        data = request.get_json()

        if not data:
            return jsonify({'error': 'Analysis data is required'}), 400

        dashboard_data = report_generator.generate_dashboard_data(data)

        return jsonify(dashboard_data)

    except Exception as e:
        logger.error(f"Error generating dashboard: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/report/generate', methods=['POST'])
@require_api_key
def generate_report():
    """Generate comprehensive Word document report"""
    try:
        data = request.get_json()

        if not data:
            return jsonify({'error': 'Analysis data is required'}), 400

        report_data = report_generator.generate_comprehensive_report(data)

        # Create response with Word document
        response = send_file(
            io.BytesIO(report_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"MA_Analysis_Report_{data.get('target_symbol', 'Unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )

        return response

    except Exception as e:
        logger.error(f"Error generating report: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/report/summary', methods=['POST'])
@require_api_key
def generate_summary_report():
    """Generate executive summary report"""
    try:
        data = request.get_json()

        if not data:
            return jsonify({'error': 'Analysis data is required'}), 400

        # Generate simplified summary report
        summary_data = report_generator._generate_summary_metrics(data)

        return jsonify({
            'summary_report': summary_data,
            'generated_at': datetime.now().isoformat()
        })

    except Exception as e:
        logger.error(f"Error generating summary report: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/charts/generate', methods=['POST'])
@require_api_key
def generate_charts():
    """Generate chart data for visualization"""
    try:
        data = request.get_json()

        if not data:
            return jsonify({'error': 'Analysis data is required'}), 400

        chart_data = report_generator._generate_chart_data(data)

        return jsonify(chart_data)

    except Exception as e:
        logger.error(f"Error generating charts: {e}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.getenv('PORT', 8080))
    app.run(host='0.0.0.0', port=port, debug=False)
