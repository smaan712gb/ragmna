"""
Full Workflow Test with Real Report Generation
Uses existing services + generates actual Excel/PDF reports for validation
Tests: MSFT → HOOD acquisition analysis
"""

import os
import sys
import json
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

# Add paths
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Test configuration
ACQUIRER = 'MSFT'
TARGET = 'HOOD'
TEST_NAME = f'FULL_WORKFLOW_{ACQUIRER}_{TARGET}_{datetime.now().strftime("%Y%m%d_%H%M%S")}'

print("="*80)
print(f"🚀 FULL WORKFLOW TEST WITH REPORT GENERATION")
print("="*80)
print(f"Test: {ACQUIRER} → {TARGET}")
print(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
print("="*80)

# Import existing services
print("\n📦 Importing services...")
sys.path.insert(0, 'services/data-ingestion')
sys.path.insert(0, 'services/llm-orchestrator')
sys.path.insert(0, 'services/three-statement-modeler')
sys.path.insert(0, 'services/dcf-valuation')

from services.data_ingestion.main import data_ingestion
from services.llm_orchestrator.main import orchestrator
from services.three_statement_modeler.main import modeler
from services.dcf_valuation.main import dcf_engine

print("✅ Services imported\n")

# Results storage
results = {
    'test_name': TEST_NAME,
    'acquirer': ACQUIRER,
    'target': TARGET,
    'timestamp': datetime.now().isoformat(),
    'phases': {}
}

# PHASE 1: Data Ingestion with Real APIs
print("-"*80)
print("PHASE 1-2: DATA INGESTION & RAG")
print("-"*80)

print(f"\n📥 Fetching comprehensive data for {TARGET}...")
target_data = data_ingestion.fetch_company_data(
    symbol=TARGET,
    data_sources=['sec_filings', 'analyst_reports', 'news']
)

print(f"✅ Data fetched")
print(f"   Status: {target_data.get('status')}")
print(f"   Company: {target_data.get('company_info', {}).get('companyName', 'Unknown')}")
print(f"   Market Cap: ${target_data.get('company_info', {}).get('mktCap', 0):,.0f}")
print(f"   Documents vectorized: {target_data.get('vectorization_results', {}).get('total_documents', 0)}")

results['phases']['data_ingestion'] = {
    'status': target_data.get('status'),
    'documents_processed': target_data.get('vectorization_results', {}).get('total_documents', 0)
}

# PHASE 3: Classification
print("\n" + "-"*80)
print("PHASE 3: CLASSIFIC ATION")
print("-"*80)

print(f"\n🏷️ Classifying {TARGET}...")
company_info = target_data.get('company_info', {})

# Synchronous classification call
import asyncio
classification = asyncio.run(
    orchestrator.classifier.classify_company_profile(TARGET, company_info)
)

print(f"✅ Classification complete")
print(f"   Classification result available")

results['phases']['classification'] = {
    'status': 'complete',
    'classification': str(classification.get('classification', 'Unknown'))[:100]
}

# PHASE 5: 3-Statement Model
print("\n" + "-"*80)
print("PHASE 5: 3-STATEMENT MODELING")
print("-"*80)

print(f"\n📊 Building 3-Statement Model...")
financial_model = modeler.generate_three_statement_model(
    company_data=target_data,
    classification=classification,
    projection_years=5
)

print(f"✅ 3SM generated")
print(f"   Projection years: {financial_model.get('projection_years', 0)}")
print(f"   Classification: {financial_model.get('classification', 'Unknown')}")

# Show projections
income_stmt = financial_model.get('income_statement', [])
if income_stmt:
    print(f"\n   📈 Revenue Projections:")
    for year_data in income_stmt[:5]:
        year = year_data.get('year', 0)
        revenue = year_data.get('revenue', 0)
        margin = year_data.get('operating_margin', 0)
        print(f"      Year {year}: ${revenue:,.0f} (Margin: {margin:.1%})")

results['phases']['3sm'] = {
    'status': 'complete',
    'projection_years': financial_model.get('projection_years', 0)
}

# PHASE 6: DCF Valuation
print("\n" + "-"*80)
print("PHASE 6: DCF VALUATION")
print("-"*80)

print(f"\n💰 Performing DCF Analysis...")
dcf_result = dcf_engine.perform_dcf_analysis(
    company_data=target_data,
    financial_model=financial_model,
    classification=classification
)

final_val = dcf_result.get('final_valuation', {})
print(f"✅ DCF complete")
print(f"   Enterprise Value: ${final_val.get('enterprise_value', 0):,.0f}")
print(f"   Equity Value: ${final_val.get('equity_value', 0):,.0f}")
print(f"   Price per Share: ${final_val.get('equity_value_per_share', 0):,.2f}")
print(f"   Current Price: ${final_val.get('current_market_price', 0):,.2f}")
print(f"   Premium/Discount: {final_val.get('premium_discount', 0):.1%}")
print(f"   WACC: {dcf_result.get('wacc', 0):.2%}")

results['phases']['dcf'] = {
    'status': 'complete',
    'enterprise_value': final_val.get('enterprise_value', 0),
    'price_per_share': final_val.get('equity_value_per_share', 0)
}

# GENERATE REPORTS
print("\n" + "="*80)
print("📄 GENERATING VALIDATION REPORTS")
print("="*80)

# Report 1: Excel Summary
print(f"\n📊 Creating Excel Summary Report...")
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Analysis Summary"
    
    # Header
    ws['A1'] = f"{ACQUIRER} → {TARGET} M&A Analysis"
    ws['A1'].font = Font(size=16, bold=True)
    ws['A2'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    
    # Company Info
    ws['A4'] = "TARGET COMPANY"
    ws['A4'].font = Font(bold=True)
    ws['A5'] = "Company Name:"
    ws['B5'] = company_info.get('companyName', 'Unknown')
    ws['A6'] = "Symbol:"
    ws['B6'] = TARGET
    ws['A7'] = "Market Cap:"
    ws['B7'] = f"${company_info.get('mktCap', 0):,.0f}"
    ws['A8'] = "Industry:"
    ws['B8'] = company_info.get('industry', 'Unknown')
    
    # Classification
    ws['A10'] = "CLASSIFICATION"
    ws['A10'].font = Font(bold=True)
    ws['A11'] = "Growth Profile:"
    ws['B11'] = classification.get('classification', 'Unknown')
    
    # 3SM Summary
    ws['A13'] = "3-STATEMENT MODEL PROJECTIONS"
    ws['A13'].font = Font(bold=True)
    ws['A14'] = "Year"
    ws['B14'] = "Revenue"
    ws['C14'] = "Net Income"
    ws['D14'] = "Operating Margin"
    
    for i, year_data in enumerate(income_stmt[:5], start=15):
        ws[f'A{i}'] = year_data.get('year', 0)
        ws[f'B{i}'] = year_data.get('revenue', 0)
        ws[f'C{i}'] = year_data.get('net_income', 0)
        ws[f'D{i}'] = year_data.get('operating_margin', 0)
    
    # DCF Valuation
    ws['A21'] = "DCF VALUATION"
    ws['A21'].font = Font(bold=True)
    ws['A22'] = "Enterprise Value:"
    ws['B22'] = final_val.get('enterprise_value', 0)
    ws['A23'] = "Equity Value:"
    ws['B23'] = final_val.get('equity_value', 0)  
    ws['A24'] = "Price per Share:"
    ws['B24'] = final_val.get('equity_value_per_share', 0)
    ws['A25'] = "Current Price:"
    ws['B25'] = final_val.get('current_market_price', 0)
    ws['A26'] = "Premium/Discount:"
    ws['B26'] = final_val.get('premium_discount', 0)
    
    excel_file = f'{TEST_NAME}_summary.xlsx'
    wb.save(excel_file)
    print(f"✅ Excel report saved: {excel_file}")
    
except Exception as e:
    print(f"⚠️ Excel generation error: {e}")
    excel_file = None

# Report 2: Word Document
print(f"\n📝 Creating Word Document Report...")
try:
    doc = Document()
    
    # Title
    title = doc.add_heading(f'{ACQUIRER} → {TARGET} M&A Analysis', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(f'Target Company: {company_info.get("companyName", "Unknown")} ({TARGET})')
    doc.add_paragraph(f'Market Capitalization: ${company_info.get("mktCap", 0):,.0f}')
    doc.add_paragraph(f'Industry: {company_info.get("industry", "Unknown")}')
    doc.add_paragraph()
    
    # Classification
    doc.add_heading('Company Classification', 2)
    doc.add_paragraph(f'Classification: {classification.get("classification", "Unknown")}')
    doc.add_paragraph()
    
    # Financial Projections
    doc.add_heading('Financial Projections (5-Year)', 2)
    doc.add_paragraph('Income Statement Forecast:')
    
    for year_data in income_stmt[:5]:
        year = year_data.get('year', 0)
        revenue = year_data.get('revenue', 0)
        ni = year_data.get('net_income', 0)
        doc.add_paragraph(
            f'Year {year}: Revenue ${revenue:,.0f}, Net Income ${ni:,.0f}',
            style='List Bullet'
        )
    
    doc.add_paragraph()
    
    # Valuation
    doc.add_heading('DCF Valuation', 2)
    doc.add_paragraph(f'Enterprise Value: ${final_val.get("enterprise_value", 0):,.0f}')
    doc.add_paragraph(f'Equity Value: ${final_val.get("equity_value", 0):,.0f}')
    doc.add_paragraph(f'Implied Price per Share: ${final_val.get("equity_value_per_share", 0):,.2f}')
    doc.add_paragraph(f'Current Market Price: ${final_val.get("current_market_price", 0):,.2f}')
    doc.add_paragraph(f'Implied Premium/(Discount): {final_val.get("premium_discount", 0):.1%}')
    
    word_file = f'{TEST_NAME}_report.docx'
    doc.save(word_file)
    print(f"✅ Word report saved: {word_file}")
    
except Exception as e:
    print(f"⚠️ Word generation error: {e}")
    word_file = None

# Report 3: JSON Data Export
print(f"\n💾 Creating JSON Data Export...")
json_export = {
    'test_name': TEST_NAME,
    'acquirer': ACQUIRER,
    'target': TARGET,
    'timestamp': datetime.now().isoformat(),
    'company_info': company_info,
    'classification': classification,
    'financial_model': {
        'income_statement': income_stmt,
        'projection_years': financial_model.get('projection_years', 0)
    },
    'dcf_valuation': {
        'enterprise_value': final_val.get('enterprise_value', 0),
        'equity_value': final_val.get('equity_value', 0),
        'price_per_share': final_val.get('equity_value_per_share', 0),
        'wacc': dcf_result.get('wacc', 0)
    },
    'results': results
}

json_file = f'{TEST_NAME}_data.json'
with open(json_file, 'w') as f:
    json.dump(json_export, f, indent=2)

print(f"✅ JSON data saved: {json_file}")

# SUMMARY
print("\n" + "="*80)
print("✅ TEST COMPLETE - REPORTS GENERATED")
print("="*80)

print(f"\n📁 Generated Files for Validation:")
if excel_file:
    print(f"   ✅ Excel: {excel_file}")
if word_file:
    print(f"   ✅ Word: {word_file}")
print(f"   ✅ JSON: {json_file}")

print(f"\n📊 Key Results to Validate:")
print(f"   • Company: {company_info.get('companyName', 'Unknown')}")
print(f"   • Classification: {classification.get('classification', 'Unknown')}")
print(f"   • 3SM Projections: {len(income_stmt)} years")
print(f"   • DCF Value: ${final_val.get('equity_value_per_share', 0):,.2f}/share")
print(f"   • Current Price: ${final_val.get('current_market_price', 0):,.2f}/share")
print(f"   • Implied Premium: {final_val.get('premium_discount', 0):.1%}")

print(f"\n🎯 Validation Checklist:")
print(f"   [ ] Open Excel file and verify data accuracy")
print(f"   [ ] Review Word document for completeness")
print(f"   [ ] Check JSON for data integrity")
print(f"   [ ] Confirm 3SM balance sheet balances")
print(f"   [ ] Validate DCF calculations")
print(f"   [ ] Review classification reasonableness")

print(f"\n✨ Test complete! Review generated files.")
